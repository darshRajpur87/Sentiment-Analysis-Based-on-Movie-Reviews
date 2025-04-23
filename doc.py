from docx import Document

# Create a new Word document
doc = Document()
doc.add_heading('Table of Contents - DishDash App', level=1)

# Add a table with 11 rows (including header) and 3 columns
table = doc.add_table(rows=11, cols=3)
table.style = 'Table Grid'

# Define headers and content
headers = ['S.No', 'Section Title', 'Page No.']
content = [
    ['I', 'Market Research on DishDash App', '3'],
    ['II', 'User Research', '8'],
    ['III', 'User Persona', '10'],
    ['IV', 'Use Case Diagram', '13'],
    ['V', 'Activity Diagram', '14'],
    ['VI', 'Sequence Diagram', '20'],
    ['VII', 'Low Fidelity Wireframes', '21'],
    ['VIII', 'Medium Fidelity Wireframes', '22'],
    ['IX', 'High Fidelity Wireframes', '24'],
    ['X', 'Prototyping', '26'],
]

# Fill in the header row with bold text
for i in range(3):
    cell = table.cell(0, i)
    cell.text = headers[i]
    run = cell.paragraphs[0].runs[0]
    run.bold = True

# Fill in the rest of the table
for row_idx, row_data in enumerate(content, start=1):
    for col_idx, item in enumerate(row_data):
        table.cell(row_idx, col_idx).text = item

# Save the document
file_path = "E:\AIML Miniprojects\Table.doc"
doc.save(file_path)

file_path
